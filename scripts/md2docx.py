#!/usr/bin/env python3
"""
Markdown to DOCX Converter
Converts markdown files to Word documents with proper formatting
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import sys


class MarkdownToDocxConverter:
    def __init__(self, md_content):
        self.content = md_content
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Setup document styles"""
        # Normal style
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

    def _add_heading(self, text, level):
        """Add heading with proper formatting"""
        heading = self.doc.add_heading(text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _add_paragraph(self, text, bold=False, italic=False):
        """Add paragraph with optional formatting"""
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return para

    def _parse_inline_formatting(self, text):
        """Parse inline formatting (bold, italic)"""
        # This is a simplified version - for full parsing, use markdown library
        # For now, return text as-is (we'll handle basic cases)
        return text

    def _is_table_line(self, line):
        """Check if line is a table separator"""
        return re.match(r'^[\|\-\+\: ]+$', line.strip())

    def _parse_table_line(self, line):
        """Parse a table row"""
        # Remove leading/trailing pipes and split
        content = line.strip()
        if content.startswith('|'):
            content = content[1:]
        if content.endswith('|'):
            content = content[:-1]
        cells = [cell.strip() for cell in content.split('|')]
        return cells

    def convert(self):
        """Convert markdown to docx"""
        lines = self.content.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i]

            # Empty lines
            if not line.strip():
                i += 1
                continue

            # Headings
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                self._add_heading(text, level)
                i += 1
                continue

            # Horizontal rule
            if re.match(r'^[-*_]{3,}$', line.strip()):
                self.doc.add_paragraph('_' * 50)
                i += 1
                continue

            # Tables
            if '|' in line and not line.startswith('```'):
                # Check if this is a table
                next_line = lines[i + 1] if i + 1 < len(lines) else ''
                if self._is_table_line(next_line):
                    # Parse header
                    header_cells = self._parse_table_line(line)
                    num_cols = len(header_cells)

                    if num_cols > 0:
                        # Skip header separator
                        i += 2

                        # Create table
                        table = self.doc.add_table(rows=1, cols=num_cols)
                        table.style = 'Table Grid'

                        # Add header row
                        for j, cell in enumerate(header_cells):
                            if j < num_cols:
                                table.rows[0].cells[j].text = cell
                                # Bold header
                                for paragraph in table.rows[0].cells[j].paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True

                        # Add data rows
                        while i < len(lines) and '|' in lines[i] and not self._is_table_line(lines[i]):
                            row_cells = self._parse_table_line(lines[i])
                            row = table.add_row()
                            for j, cell in enumerate(row_cells):
                                if j < len(row.cells):
                                    row.cells[j].text = cell
                            i += 1

                    # Add spacing after table
                    self.doc.add_paragraph()
                    continue

            # Code blocks
            if line.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                i += 1  # Skip closing ```

                # Add code as a paragraph with monospace-ish font
                para = self.doc.add_paragraph('\n'.join(code_lines))
                para.style = 'No Spacing'
                for run in para.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(64, 64, 64)

                continue

            # Lists
            if re.match(r'^\s*[-*+]\s+', line):
                # Unordered list
                list_items = []
                while i < len(lines) and re.match(r'^\s*[-*+]\s+', lines[i]):
                    item = re.sub(r'^\s*[-*+]\s+', '', lines[i])
                    list_items.append(item)
                    i += 1

                for item in list_items:
                    para = self.doc.add_paragraph(item, style='List Bullet')

                continue

            if re.match(r'^\s*\d+\.\s+', line):
                # Ordered list
                list_items = []
                while i < len(lines) and re.match(r'^\s*\d+\.\s+', lines[i]):
                    item = re.sub(r'^\s*\d+\.\s+', '', lines[i])
                    list_items.append(item)
                    i += 1

                for item in list_items:
                    para = self.doc.add_paragraph(item, style='List Number')

                continue

            # Blockquotes
            if line.startswith('>'):
                quote_text = re.sub(r'^>\s*', '', line)
                para = self.doc.add_paragraph(quote_text)
                para.style = 'Quote'
                i += 1
                continue

            # Regular paragraph
            self._add_paragraph(line.strip())
            i += 1

        return self.doc

    def save(self, output_path):
        """Save the document"""
        self.doc.save(output_path)
        return output_path


def main():
    if len(sys.argv) < 3:
        print("Usage: python md2docx.py <input.md> <output.docx>")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    # Read markdown file
    with open(input_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Convert to docx
    converter = MarkdownToDocxConverter(md_content)
    converter.convert()

    # Save
    converter.save(output_file)
    print(f"✓ Converted {input_file} to {output_file}")


if __name__ == '__main__':
    main()
