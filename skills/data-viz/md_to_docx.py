#!/usr/bin/env python3
"""
Convert Markdown to DOCX with basic formatting
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import sys

def parse_md_to_docx(md_file, docx_file):
    doc = Document()

    # Add title
    doc.add_heading('AI Data Center Analysis - Africa', 0)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    current_list = None
    list_items = []

    for line in lines:
        line = line.rstrip('\n')

        # Skip empty lines
        if not line.strip():
            if current_list:
                for item in list_items:
                    p = doc.add_paragraph(item, style='List Bullet')
                list_items = []
                current_list = None
            continue

        # Headers
        if line.startswith('# '):
            if current_list:
                for item in list_items:
                    doc.add_paragraph(item, style='List Bullet')
                list_items = []
                current_list = None
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            if current_list:
                for item in list_items:
                    doc.add_paragraph(item, style='List Bullet')
                list_items = []
                current_list = None
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            if current_list:
                for item in list_items:
                    doc.add_paragraph(item, style='List Bullet')
                list_items = []
                current_list = None
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            if current_list:
                for item in list_items:
                    doc.add_paragraph(item, style='List Bullet')
                list_items = []
                current_list = None
            doc.add_heading(line[5:], level=4)

        # Bold text **text**
        elif '**' in line:
            if current_list:
                list_items.append(line)
            else:
                p = doc.add_paragraph()
                parts = re.split(r'\*\*(.+?)\*\*', line)
                for i, part in enumerate(parts):
                    if i % 2 == 1:  # Bold
                        run = p.add_run(part)
                        run.bold = True
                    else:
                        if part.strip():
                            p.add_run(part)

        # Lists (bullets with - or *)
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            current_list = 'bullet'
            list_items.append(line.strip()[2:])

        # Numbered lists
        elif line.strip().startswith(('1. ', '2. ', '3. ', '4. ', '5. ',
                                   '6. ', '7. ', '8. ', '9. ', '10. ')):
            if current_list and current_list != 'numbered':
                for item in list_items:
                    doc.add_paragraph(item, style='List Bullet')
                list_items = []
            current_list = 'numbered'
            list_items.append(line.strip()[3:])

        # Tables (simple CSV-like format)
        elif line.startswith('|'):
            # Skip separator lines
            if '|-' in line or '---' in line:
                continue

            parts = [p.strip() for p in line.split('|')[1:-1]]
            if parts and len(parts) > 1:
                # Check if it's a header row
                is_header = any('Rank' in p or 'Location' in p or 'Country' in p for p in parts)

                if is_header:
                    table = doc.add_table(rows=1, cols=len(parts))
                    table.style = 'Light Grid Accent 1'
                    for i, part in enumerate(parts):
                        table.rows[0].cells[i].text = part
                        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
                else:
                    # Add row to existing table
                    try:
                        table.add_row()
                        for i, part in enumerate(parts):
                            table.rows[-1].cells[i].text = part
                    except:
                        pass

        # Regular paragraph
        else:
            if current_list:
                for item in list_items:
                    doc.add_paragraph(item, style='List Bullet')
                list_items = []
                current_list = None
            if line.strip():
                p = doc.add_paragraph(line)

    doc.save(docx_file)
    print(f"✓ Converted {md_file} to {docx_file}")

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Usage: python md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)

    parse_md_to_docx(sys.argv[1], sys.argv[2])
